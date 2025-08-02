from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_from_directory, session
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta, timezone
import calendar
import os
import pandas as pd
from werkzeug.utils import secure_filename
import numpy as np
import re
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from sqlalchemy import text
import secrets
import threading
import time

app = Flask(__name__)
app.config['SECRET_KEY'] = 'tu_clave_secreta_aqui'
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///aseo.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Configurar ruta estática adicional para las imágenes
app.static_folder = 'static'
app.static_url_path = '/static'

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads', 'salud_casanare')
ALLOWED_EXTENSIONS = {'xlsx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_expired_users():
    """Elimina usuarios temporales que han expirado"""
    with app.app_context():
        expired_users = User.query.filter(
            User.is_temporary == True,
            User.expires_at < datetime.now(timezone.utc)
        ).all()
        
        for user in expired_users:
            db.session.delete(user)
        
        if expired_users:
            db.session.commit()
            print(f"Eliminados {len(expired_users)} usuarios temporales expirados")

# Modelos
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), default='user')
    created_at = db.Column(db.DateTime, default=datetime.now(timezone.utc))
    is_temporary = db.Column(db.Boolean, default=False)
    expires_at = db.Column(db.DateTime, nullable=True)
    temp_password = db.Column(db.String(120), nullable=True)
    login_token = db.Column(db.String(100), unique=True, nullable=True)
    login_token_used = db.Column(db.Boolean, default=False)

class Task(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text)
    status = db.Column(db.String(20), default='pending')  # pending, in_progress, completed
    assigned_to = db.Column(db.Integer, db.ForeignKey('user.id'))
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'))
    created_at = db.Column(db.DateTime, default=datetime.now(timezone.utc))
    completed_at = db.Column(db.DateTime)
    scheduled_for = db.Column(db.DateTime)

class ActivityLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey('task.id'))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    action = db.Column(db.String(50), nullable=False)
    details = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now(timezone.utc))

class Planilla(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    mes = db.Column(db.Integer, nullable=False)  # 1-12 para los meses
    año = db.Column(db.Integer, nullable=False)
    titulo = db.Column(db.String(200), nullable=False)
    descripcion = db.Column(db.Text)
    url_google_drive = db.Column(db.String(500), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.now(timezone.utc))
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'))

# Modelo para mensajes importantes
class ImportantMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.Text, nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=datetime.now(timezone.utc), onupdate=datetime.now(timezone.utc))

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def init_db():
    with app.app_context():
        db.create_all()
        
        # Migración: agregar campos de usuario temporal si no existen
        try:
            # Verificar si las columnas ya existen
            inspector = db.inspect(db.engine)
            existing_columns = [col['name'] for col in inspector.get_columns('user')]
            
            if 'is_temporary' not in existing_columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE user ADD COLUMN is_temporary BOOLEAN DEFAULT FALSE'))
                    conn.commit()
                print("Columna is_temporary agregada")
                
            if 'expires_at' not in existing_columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE user ADD COLUMN expires_at DATETIME'))
                    conn.commit()
                print("Columna expires_at agregada")
                
            if 'temp_password' not in existing_columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE user ADD COLUMN temp_password VARCHAR(120)'))
                    conn.commit()
                print("Columna temp_password agregada")
                
            if 'login_token' not in existing_columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE user ADD COLUMN login_token VARCHAR(100)'))
                    conn.commit()
                print("Columna login_token agregada")

            if 'login_token_used' not in existing_columns:
                with db.engine.connect() as conn:
                    conn.execute(text('ALTER TABLE user ADD COLUMN login_token_used BOOLEAN DEFAULT FALSE'))
                    conn.commit()
                print("Columna login_token_used agregada")
                
        except Exception as e:
            print(f"Error en migración: {e}")
        
        # Crear usuario admin por defecto si no existe
        if not User.query.filter_by(username='admin').first():
            admin_user = User(
                username='admin',
                password_hash=generate_password_hash('admin123'),
                role='admin'
            )
            db.session.add(admin_user)
            db.session.commit()
            print("Usuario admin creado por defecto")

        # Crear usuarios por defecto si no existen
        if not User.query.filter_by(username='root').first():
            root_user = User(
                username='root',
                password_hash=generate_password_hash('aseo2025slclabor'),
                role='admin'
            )
            db.session.add(root_user)
        
        if not User.query.filter_by(username='julio').first():
            julio_user = User(
                username='julio',
                password_hash=generate_password_hash('julio21200521A'),
                role='user'
            )
            db.session.add(julio_user)
        
        # Crear planilla de enero 2025 si no existe
        if not Planilla.query.filter_by(mes=1, año=2025).first():
            planilla_enero = Planilla(
                mes=1,
                año=2025,
                titulo='Planilla de Control - Enero 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1PidVqTqWdb_1C1t43iOhz_YCr8JVFSK9/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_enero)

        # Crear planilla de febrero 2025 si no existe
        if not Planilla.query.filter_by(mes=2, año=2025).first():
            planilla_febrero = Planilla(
                mes=2,
                año=2025,
                titulo='Planilla de Control - Febrero 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1MjGp8u1g_TbdG4PIROhjUFT_gWIHq_iB/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_febrero)

        # Crear planilla de marzo 2025 si no existe
        if not Planilla.query.filter_by(mes=3, año=2025).first():
            planilla_marzo = Planilla(
                mes=3,
                año=2025,
                titulo='Planilla de Control - Marzo 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1RrD3Dtk4nPI65WDr1JmLUjLNm-W0pbqu/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_marzo)

        # Crear planilla de abril 2025 si no existe
        if not Planilla.query.filter_by(mes=4, año=2025).first():
            planilla_abril = Planilla(
                mes=4,
                año=2025,
                titulo='Planilla de Control - Abril 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1BMPMShT5PENhwy9g5LePGXjuMywjllCB/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_abril)

        # Crear planilla de mayo 2025 si no existe
        if not Planilla.query.filter_by(mes=5, año=2025).first():
            planilla_mayo = Planilla(
                mes=5,
                año=2025,
                titulo='Planilla de Control - Mayo 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1ktZw4-fNvCKbb7alQX34Gav1n9vA0xrS/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_mayo)

        # Crear planilla de junio 2025 si no existe
        if not Planilla.query.filter_by(mes=6, año=2025).first():
            planilla_junio = Planilla(
                mes=6,
                año=2025,
                titulo='Planilla de Control - Junio 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1TSKZEwXONCQYCNkPVWPU3gGzGeOGtRhS/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_junio)

        # Crear planilla de julio 2025 si no existe
        if not Planilla.query.filter_by(mes=7, año=2025).first():
            planilla_julio = Planilla(
                mes=7,
                año=2025,
                titulo='Planilla de Control - Julio 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1EHOT79WLtniLa6nECipvdT0G7QT-BPpA/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_julio)

        # Crear planilla de agosto 2025 si no existe
        if not Planilla.query.filter_by(mes=8, año=2025).first():
            planilla_agosto = Planilla(
                mes=8,
                año=2025,
                titulo='Planilla de Control - Agosto 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/10Sgg-HUKumAHZbze7mkVif3pl2EUB4Rv/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_agosto)

        # Crear planilla de septiembre 2025 si no existe
        if not Planilla.query.filter_by(mes=9, año=2025).first():
            planilla_septiembre = Planilla(
                mes=9,
                año=2025,
                titulo='Planilla de Control - Septiembre 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1ohuWRs15vq4spKhVDLjgbQW0xMKPIzjb/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_septiembre)

        # Crear planilla de octubre 2025 si no existe
        if not Planilla.query.filter_by(mes=10, año=2025).first():
            planilla_octubre = Planilla(
                mes=10,
                año=2025,
                titulo='Planilla de Control - Octubre 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1HQUmzc5DBo3vpUusq6gFh9JMAV383e96/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_octubre)

        # Crear planilla de noviembre 2025 si no existe
        if not Planilla.query.filter_by(mes=11, año=2025).first():
            planilla_noviembre = Planilla(
                mes=11,
                año=2025,
                titulo='Planilla de Control - Noviembre 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1wFRka_DdNv1_UUcnGsplPP8wO0U7CF9R/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_noviembre)

        # Crear planilla de diciembre 2025 si no existe
        if not Planilla.query.filter_by(mes=12, año=2025).first():
            planilla_diciembre = Planilla(
                mes=12,
                año=2025,
                titulo='Planilla de Control - Diciembre 2025',
                descripcion='Planilla de control de servicios de aseo para Salud Casanare',
                url_google_drive='https://docs.google.com/spreadsheets/d/1fzYNqB9yrHTh1bWfnLtrg1--FT4Nt5T2/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
                created_by=1  # ID del usuario root
            )
            db.session.add(planilla_diciembre)
        
        db.session.commit()

        # Crear algunas tareas de ejemplo
        if not Task.query.first():
            tasks = [
                {
                    'title': 'Limpieza de Oficinas',
                    'description': 'Limpieza general de todas las oficinas del primer piso',
                    'status': 'completed',
                    'scheduled_for': datetime.now(timezone.utc)
                },
                {
                    'title': 'Desinfección de Áreas Comunes',
                    'description': 'Desinfección de áreas comunes y pasillos',
                    'status': 'completed',
                    'scheduled_for': datetime.now(timezone.utc)
                },
                {
                    'title': 'Limpieza de Ventanas',
                    'description': 'Limpieza de ventanas exteriores',
                    'status': 'pending',
                    'scheduled_for': datetime.now(timezone.utc)
                }
            ]
            
            for task_data in tasks:
                task = Task(**task_data)
                db.session.add(task)
            
            db.session.commit()

        # Crear mensaje importante por defecto si no existe
        if not ImportantMessage.query.first():
            msg = ImportantMessage(content='', is_active=False)
            db.session.add(msg)
            db.session.commit()

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/inicio')
def inicio():
    meses = []
    nombres_meses = {
        1: 'Enero',
        2: 'Febrero',
        3: 'Marzo',
        4: 'Abril',
        5: 'Mayo',
        6: 'Junio',
        7: 'Julio',
        8: 'Agosto',
        9: 'Septiembre',
        10: 'Octubre',
        11: 'Noviembre',
        12: 'Diciembre'
    }
    for mes in range(1, 13):
        meses.append({
            'numero': mes,
            'nombre': nombres_meses[mes],
            'imagen': f'static/images/meses/{mes}.jpg'
        })
    return render_template('inicio.html', meses=meses)

@app.route('/mes/<int:mes>')
@login_required
def ver_mes(mes):
    if 1 <= mes <= 12:
        # Nombres de los meses en español
        nombres_meses = {
            1: 'Enero',
            2: 'Febrero',
            3: 'Marzo',
            4: 'Abril',
            5: 'Mayo',
            6: 'Junio',
            7: 'Julio',
            8: 'Agosto',
            9: 'Septiembre',
            10: 'Octubre',
            11: 'Noviembre',
            12: 'Diciembre'
        }

        # Tarjetas estándar para todos los meses
        tarjetas_estandar = [
            {
                'titulo': 'Salud Casanare',
                'descripcion': 'Servicios especializados de limpieza y desinfección para el sector salud',
                'imagen': 'salud casanare.png',
                'icono': 'fa-hospital'
            },
            {
                'titulo': 'Laboratorio',
                'descripcion': 'Mantenimiento y limpieza de áreas críticas y equipos especializados',
                'imagen': 'laboratorio.jpg',
                'icono': 'fa-flask'
            }
        ]
        
        # Crear la información del mes actual
        mes_actual = {
            'nombre': nombres_meses[mes],
            'tarjetas': tarjetas_estandar
        }
        
        return render_template('mes.html', 
                            mes=mes, 
                            nombre_mes=mes_actual['nombre'],
                            tarjetas=mes_actual['tarjetas'])
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash('Usuario o contraseña incorrectos', 'error')
    
    return render_template('login.html')

@app.route('/dashboard')
@login_required
def dashboard():
    # Obtener nombres de los meses
    nombres_meses = [
        'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
        'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre'
    ]
    # Listar imágenes reales en la carpeta
    ruta_img = os.path.join(app.static_folder, 'images', 'meses')
    imagenes = [f for f in os.listdir(ruta_img) if f.lower().endswith(('.jpg','.jpeg','.png','.webp'))]
    imagenes.sort()  # Orden alfabético
    # Emparejar imágenes con meses para 2025
    meses = []
    for i, img in enumerate(imagenes):
        if i < len(nombres_meses):
            meses.append({'nombre': nombres_meses[i], 'archivo': img})
    # Emparejar imágenes con meses para 2026 (idéntico a 2025)
    meses_2026 = []
    for i, img in enumerate(imagenes):
        if i < len(nombres_meses):
            meses_2026.append({'nombre': nombres_meses[i], 'archivo': img})
    # Obtener estadísticas
    stats = {
        'pending_tasks': Task.query.filter_by(status='pending').count(),
        'completed_tasks': Task.query.filter_by(status='completed').count(),
        'active_users': User.query.filter(User.role == 'user').count(),
        'avg_completion_time': '24h'  # Esto debería calcularse basado en datos reales
    }
    # Obtener actividades recientes
    recent_activities = ActivityLog.query.order_by(ActivityLog.created_at.desc()).limit(3).all()
    # Obtener próximas tareas
    upcoming_tasks = Task.query.filter_by(status='pending').order_by(Task.scheduled_for.asc()).limit(3).all()
    
    # Año dinámico para la sección 2026 (puede ser configurable)
    year_2026 = session.get('global_year', 2026)  # Usar el año de la sesión o 2026 por defecto

    # Estado de la sección 2025
    section_2025_enabled = session.get('section_2025_enabled', True)
    section_2025_message = session.get('section_2025_message', 'Sección temporalmente no disponible. Posible actualización de plataforma en curso. Por favor, contacte al administrador para más información.')

    important_msg = ImportantMessage.query.first()
    return render_template('dashboard.html', 
                         stats=stats,
                         recent_activities=recent_activities,
                         upcoming_tasks=upcoming_tasks,
                         meses=meses,
                         meses_2026=meses_2026,
                         year_2026=year_2026,
                         section_2025_enabled=section_2025_enabled,
                         section_2025_message=section_2025_message,
                         important_msg=important_msg)

@app.route('/logout')
@login_required
def logout():
    # Limpiar sesión de admin también
    session.pop('admin_authenticated', None)
    session.pop('admin_login_time', None)
    
    logout_user()
    flash('Has cerrado sesión exitosamente.', 'info')
    return redirect(url_for('login'))

# API endpoints para actualizar datos en tiempo real
@app.route('/api/stats')
@login_required
def get_stats():
    stats = {
        'pending_tasks': Task.query.filter_by(status='pending').count(),
        'completed_tasks': Task.query.filter_by(status='completed').count(),
        'active_users': User.query.filter(User.role == 'user').count(),
        'avg_completion_time': '24h'
    }
    return jsonify(stats)

@app.route('/imagenes/<path:filename>')
def serve_images(filename):
    return send_from_directory('imagenes salud casanare laboratorio', filename)

@app.route('/salud_casanare_google/<int:mes>')
@login_required
def salud_casanare_google(mes):
    # Verificar si la sección 2025 está desactivada
    section_2025_enabled = session.get('section_2025_enabled', True)
    if not section_2025_enabled:
        message = session.get('section_2025_message', 'Sección temporalmente no disponible. Posible actualización de plataforma en curso. Por favor, contacte al administrador para más información.')
        flash(message, 'warning')
        return redirect(url_for('dashboard'))
    
    # Obtener la planilla del mes específico
    planilla = Planilla.query.filter_by(mes=mes, año=2025).first()
    
    # Nombres de los meses en español
    nombres_meses = {
        1: 'Enero',
        2: 'Febrero',
        3: 'Marzo',
        4: 'Abril',
        5: 'Mayo',
        6: 'Junio',
        7: 'Julio',
        8: 'Agosto',
        9: 'Septiembre',
        10: 'Octubre',
        11: 'Noviembre',
        12: 'Diciembre'
    }
    
    return render_template('salud_casanare_google.html', 
                         planilla=planilla,
                         nombres_meses=nombres_meses,
                         mes_actual=mes)

@app.route('/descargar_planilla/<int:mes>/<formato>')
@login_required
def descargar_planilla(mes, formato):
    # Verificar si la sección 2025 está desactivada
    section_2025_enabled = session.get('section_2025_enabled', True)
    if not section_2025_enabled:
        message = session.get('section_2025_message', 'Sección temporalmente no disponible. Posible actualización de plataforma en curso. Por favor, contacte al administrador para más información.')
        flash(message, 'warning')
        return redirect(url_for('dashboard'))
    
    # Obtener la planilla del mes específico
    planilla = Planilla.query.filter_by(mes=mes, año=2025).first()
    
    if not planilla:
        flash('Planilla no encontrada', 'error')
        return redirect(url_for('dashboard'))
    
    # Nombres de los meses en español
    nombres_meses = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    nombre_mes = nombres_meses.get(mes, f'Mes_{mes}')
    nombre_archivo = f"Planilla_Salud_Casanare_{nombre_mes}_{2025}"
    
    try:
        # Extraer el ID del archivo de la URL de Google Drive
        url_pattern = r'/d/([a-zA-Z0-9-_]+)/'
        match = re.search(url_pattern, planilla.url_google_drive)
        
        if not match:
            flash('URL de Google Drive no válida', 'error')
            return redirect(url_for('salud_casanare_google', mes=mes))
        
        file_id = match.group(1)
        
        if formato == 'excel':
            # URL de descarga directa para Excel
            download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"
            response = requests.get(download_url)
            
            if response.status_code == 200:
                import tempfile
                import os
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
                temp_file.write(response.content)
                temp_file.close()
                
                return send_from_directory(
                    os.path.dirname(temp_file.name),
                    os.path.basename(temp_file.name),
                    as_attachment=True,
                    download_name=f"{nombre_archivo}.xlsx"
                )
            else:
                flash('Error al descargar el archivo Excel', 'error')
                return redirect(url_for('salud_casanare_google', mes=mes))
                
        elif formato == 'pdf':
            # URL de descarga directa para PDF
            download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=pdf"
            response = requests.get(download_url)
            
            if response.status_code == 200:
                import tempfile
                import os
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                temp_file.write(response.content)
                temp_file.close()
                
                return send_from_directory(
                    os.path.dirname(temp_file.name),
                    os.path.basename(temp_file.name),
                    as_attachment=True,
                    download_name=f"{nombre_archivo}.pdf"
                )
            else:
                flash('Error al descargar el archivo PDF', 'error')
                return redirect(url_for('salud_casanare_google', mes=mes))
                
        elif formato == 'word':
            # Descargar la planilla real de Google Drive como PDF y crear Word elegante
            try:
                # Descargar como PDF desde Google Drive para mantener formato
                pdf_download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=pdf"
                response = requests.get(pdf_download_url)
                
                if response.status_code == 200:
                    import tempfile
                    import os
                    from docx import Document
                    from docx.shared import Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.shared import OxmlElement, qn
                    
                    # Crear documento Word elegante
                    doc = Document()
                    
                    # Configurar márgenes
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                    
                    # Título principal
                    title = doc.add_heading(f'Planilla de Control - {nombre_mes} 2025', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Línea separadora
                    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Información de la empresa
                    company_heading = doc.add_heading('Grupo Servis Aseo S.L', level=1)
                    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    services_para = doc.add_paragraph('Servicios de Limpieza y Mantenimiento')
                    services_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Detalles en formato tabla
                    doc.add_paragraph('')  # Espacio
                    
                    # Crear tabla de información
                    info_table = doc.add_table(rows=3, cols=2)
                    info_table.style = 'Table Grid'
                    info_table.autofit = True
                    
                    # Información de la tabla
                    info_table.cell(0, 0).text = 'Fecha de Descarga:'
                    info_table.cell(0, 1).text = datetime.now(timezone.utc).strftime("%d/%m/%Y")
                    info_table.cell(1, 0).text = 'Mes:'
                    info_table.cell(1, 1).text = nombre_mes
                    info_table.cell(2, 0).text = 'Año:'
                    info_table.cell(2, 1).text = '2025'
                    
                    # Espacio
                    doc.add_paragraph('')
                    
                    # Nota importante
                    note_heading = doc.add_heading('Información Importante', level=2)
                    note_para = doc.add_paragraph('Esta es una versión en formato Word de la planilla original de Google Drive.')
                    note_para2 = doc.add_paragraph('Para acceder a la planilla completa con formato original, utilice la opción "Descargar PDF".')
                    
                    # Espacio
                    doc.add_paragraph('')
                    
                    # Información de acceso
                    access_heading = doc.add_heading('Acceso a la Planilla Original', level=2)
                    access_para = doc.add_paragraph('• La planilla original se encuentra en Google Drive')
                    access_para2 = doc.add_paragraph('• Para editar o ver en tiempo real, acceda directamente a Google Drive')
                    access_para3 = doc.add_paragraph('• Los datos se actualizan automáticamente en la planilla original')
                    
                    # Pie de página
                    doc.add_paragraph('')
                    footer_para = doc.add_paragraph('Documento generado automáticamente por el sistema de gestión de planillas')
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    footer_para.style = 'Intense Quote'
                    
                    # Guardar archivo Word
                    word_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    word_temp.close()
                    doc.save(word_temp.name)
                    
                    return send_from_directory(
                        os.path.dirname(word_temp.name),
                        os.path.basename(word_temp.name),
                        as_attachment=True,
                        download_name=f"{nombre_archivo}.docx"
                    )
                else:
                    flash('Error al acceder a la planilla de Google Drive', 'error')
                    return redirect(url_for('salud_casanare_google', mes=mes))
                    
            except Exception as e:
                flash(f'Error al generar el documento Word: {str(e)}', 'error')
                return redirect(url_for('salud_casanare_google', mes=mes))
            
        else:
            flash('Formato no válido', 'error')
            return redirect(url_for('salud_casanare_google', mes=mes))
            
    except Exception as e:
        flash(f'Error al descargar el archivo: {str(e)}', 'error')
        return redirect(url_for('salud_casanare_google', mes=mes))

@app.route('/admin/get_2025_status')
@login_required
def get_2025_status():
    if not check_admin_session():
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        # Obtener estado desde la sesión
        enabled = session.get('section_2025_enabled', True)  # Por defecto activada
        message = session.get('section_2025_message', 'Sección temporalmente no disponible. Posible actualización de plataforma en curso. Por favor, contacte al administrador para más información.')
        
        return jsonify({
            'success': True,
            'enabled': enabled,
            'message': message
        })
        
    except Exception as e:
        return jsonify({'error': f'Error al obtener estado: {str(e)}'}), 500

@app.route('/admin/update_2025_status', methods=['POST'])
@login_required
def update_2025_status():
    if not check_admin_session():
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        data = request.get_json()
        enabled = data.get('enabled', True)
        message = data.get('message', '')
        
        if not message.strip():
            return jsonify({'error': 'El mensaje no puede estar vacío'}), 400
        
        # Guardar en la sesión
        session['section_2025_enabled'] = enabled
        session['section_2025_message'] = message
        
        return jsonify({
            'success': True,
            'message': f'Estado de sección 2025 actualizado exitosamente',
            'enabled': enabled,
            'message': message
        })
        
    except Exception as e:
        return jsonify({'error': f'Error al actualizar estado: {str(e)}'}), 500

@app.route('/laboratorio_google/<int:mes>')
@login_required
def laboratorio_google(mes):
    # Verificar si la sección 2025 está desactivada
    section_2025_enabled = session.get('section_2025_enabled', True)
    if not section_2025_enabled:
        message = session.get('section_2025_message', 'Sección temporalmente no disponible. Posible actualización de plataforma en curso. Por favor, contacte al administrador para más información.')
        flash(message, 'warning')
        return redirect(url_for('dashboard'))
    
    # Obtener la planilla del mes específico
    planilla = Planilla.query.filter_by(mes=mes, año=2025).first()
    
    # URLs específicas para cada mes
    urls_especificas = {
        1: 'https://docs.google.com/spreadsheets/d/1dDlOInXVAW-3BeBe-_q8IePx0cM70nYk/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        2: 'https://docs.google.com/spreadsheets/d/1D1e_LlQCUz84yOH96oDFp1SHGZzTeDwi/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        3: 'https://docs.google.com/spreadsheets/d/1j45Vzid9SmBSpLD51Abtp_mMG-1zgmOb/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        4: 'https://docs.google.com/spreadsheets/d/1akbFfZqAs3eAlJ_TYesrrhM6XUN8FeXP/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        5: 'https://docs.google.com/spreadsheets/d/1_dOqGMcHK-J51C6R8zHLgjocQ-oETJAl/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        6: 'https://docs.google.com/spreadsheets/d/1282tMi37hKiyvm3nPvMs7yylBwWKryyF/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        7: 'https://docs.google.com/spreadsheets/d/1JnMj_t3idxzK2EMsVw0wJmEWt9Jf2MpK/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        8: 'https://docs.google.com/spreadsheets/d/1Fa8cKJzLVdX4VD9dynPDNnXhrMfVeNjl/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        9: 'https://docs.google.com/spreadsheets/d/1RLKSFFKISJMEL64Lh02SqYcSQsTr76Rh/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        10: 'https://docs.google.com/spreadsheets/d/1tLWPyRIJI4j78mD8Khu6u3T6aII2o49l/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        11: 'https://docs.google.com/spreadsheets/d/1oOPs3Q2Wk8a6ZDMXgqaS50meNoXgCRiy/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true',
        12: 'https://docs.google.com/spreadsheets/d/1VGL3TSMjS7qh5aCxzFPjYZUkXAPPPGhU/edit?usp=sharing&ouid=116069373546627717051&rtpof=true&sd=true'
    }
    
    # Si existe una URL específica para el mes, usarla
    if mes in urls_especificas:
        planilla.url_google_drive = urls_especificas[mes]
    
    nombres_meses = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    return render_template('laboratorio_google.html', 
                         planilla=planilla,
                         nombres_meses=nombres_meses,
                         mes_actual=mes)

# Nueva ruta para ver los meses de 2026
@app.route('/mes_2026/<int:mes>')
@login_required
def ver_mes_2026(mes):
    if 1 <= mes <= 12:
        nombres_meses = {
            1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
            5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
            9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
        }
        tarjetas_estandar = [
            {
                'titulo': 'Salud Casanare',
                'descripcion': 'Servicios especializados de limpieza y desinfección para el sector salud',
                'imagen': 'salud casanare.png',
                'icono': 'fa-hospital'
            },
            {
                'titulo': 'Laboratorio',
                'descripcion': 'Mantenimiento y limpieza de áreas críticas y equipos especializados',
                'imagen': 'laboratorio.jpg',
                'icono': 'fa-flask'
            }
        ]
        mes_actual = {
            'nombre': nombres_meses[mes],
            'tarjetas': tarjetas_estandar
        }
        
        # Obtener el año dinámico
        year_2026 = session.get('global_year', 2026)
        
        return render_template('mes_2026.html', 
                            mes=mes, 
                            nombre_mes=mes_actual['nombre'],
                            tarjetas=mes_actual['tarjetas'],
                            year_2026=year_2026)
    return redirect(url_for('dashboard'))

@app.route('/salud_casanare2026_google/<int:mes>')
@login_required
def salud_casanare2026_google(mes):
    # Obtener el año dinámico
    year_2026 = session.get('global_year', 2026)
    
    # Obtener la planilla del mes específico para el año dinámico
    planilla = Planilla.query.filter_by(mes=mes, año=year_2026).first()
    nombres_meses = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    return render_template('salud_casanare2026_google.html', 
                         planilla=planilla,
                         nombres_meses=nombres_meses,
                         mes_actual=mes,
                         year_2026=year_2026)

@app.route('/laboratorio2026_google/<int:mes>')
@login_required
def laboratorio2026_google(mes):
    # Obtener el año dinámico
    year_2026 = session.get('global_year', 2026)
    
    # Obtener la planilla del mes específico para el año dinámico
    planilla = Planilla.query.filter_by(mes=mes, año=year_2026).first()
    nombres_meses = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    return render_template('laboratorio2026_google.html', 
                         planilla=planilla,
                         nombres_meses=nombres_meses,
                         mes_actual=mes,
                         year_2026=year_2026)

@app.route('/programacion')
@login_required
def programacion():
    # Obtener fecha actual
    current_date = datetime.now(timezone.utc)
    current_month = current_date.month
    current_year = current_date.year
    
    # Nombres de los meses en español
    month_names = [
        'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
        'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre'
    ]
    
    current_month_name = month_names[current_month - 1]
    
    return render_template('programacion.html', 
                         current_month_name=current_month_name,
                         current_year=current_year)

def check_admin_session():
    """Verifica si la sesión de admin está activa y no ha expirado"""
    if not session.get('admin_authenticated'):
        return False
    
    # Verificar si la sesión de admin no ha expirado (30 minutos)
    admin_login_time = session.get('admin_login_time')
    if admin_login_time:
        try:
            login_time = datetime.fromisoformat(admin_login_time)
            if datetime.now(timezone.utc) - login_time < timedelta(minutes=30):
                return True
            else:
                # Sesión expirada, limpiar
                session.pop('admin_authenticated', None)
                session.pop('admin_login_time', None)
                return False
        except:
            # Error al parsear la fecha, limpiar sesión
            session.pop('admin_authenticated', None)
            session.pop('admin_login_time', None)
            return False
    
    return False

@app.route('/admin')
@login_required
def admin_panel():
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    
    # Limpiar usuarios expirados
    cleanup_expired_users()
    
    # Obtener estadísticas del sistema
    total_users = User.query.count()
    total_planillas = Planilla.query.count()
    total_tasks = Task.query.count()
    completed_tasks = Task.query.filter_by(status='completed').count()
    
    # Obtener usuarios recientes
    recent_users = User.query.order_by(User.created_at.desc()).limit(5).all()
    
    # Obtener usuarios temporales activos
    active_temporary_users = User.query.filter(
        User.is_temporary == True,
        User.expires_at > datetime.now(timezone.utc)
    ).order_by(User.expires_at.asc()).all()
    
    # Obtener planillas recientes
    recent_planillas = Planilla.query.order_by(Planilla.created_at.desc()).limit(5).all()
    
    # Obtener todas las tareas
    all_tasks = Task.query.order_by(Task.created_at.desc()).limit(10).all()
    
    # Obtener el año dinámico
    year_2026 = session.get('global_year', 2026)
    
    return render_template('admin.html',
                         total_users=total_users,
                         total_planillas=total_planillas,
                         total_tasks=total_tasks,
                         completed_tasks=completed_tasks,
                         recent_users=recent_users,
                         active_temporary_users=active_temporary_users,
                         recent_planillas=recent_planillas,
                         all_tasks=all_tasks,
                         year_2026=year_2026)

@app.route('/admin/users')
@login_required
def admin_users():
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    
    users = User.query.all()
    return render_template('admin_users.html', users=users)

@app.route('/admin/create_user', methods=['GET', 'POST'])
@login_required
def create_user():
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        role = request.form.get('role', 'user')
        
        if User.query.filter_by(username=username).first():
            flash('El nombre de usuario ya existe.', 'error')
        else:
            new_user = User(
                username=username,
                password_hash=generate_password_hash(password),
                role=role
            )
            db.session.add(new_user)
            db.session.commit()
            flash('Usuario creado exitosamente.', 'success')
            return redirect(url_for('admin_users'))
    
    return render_template('create_user.html')

@app.route('/admin/create_temporary_user', methods=['POST'])
@login_required
def create_temporary_user():
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return jsonify({'error': 'No autenticado como administrador'}), 401
    
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        duration = data.get('duration')  # 4, 15, 30, 60, 480, 1440 (minutos)
        
        if not username or not password or not duration:
            return jsonify({'error': 'Faltan datos requeridos'}), 400
        
        if User.query.filter_by(username=username).first():
            return jsonify({'error': 'El nombre de usuario ya existe'}), 400
        
        # Calcular tiempo de expiración
        duration_minutes = int(duration)
        expires_at = datetime.now(timezone.utc) + timedelta(minutes=duration_minutes)
        login_token = secrets.token_urlsafe(32)
        
        # Crear usuario temporal
        new_user = User(
            username=username,
            password_hash=generate_password_hash(password),
            role='user', # Siempre se crea como 'user', nunca como 'admin'
            is_temporary=True,
            expires_at=expires_at,
            temp_password=password,
            login_token=login_token
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        # Limpiar usuarios expirados
        cleanup_expired_users()
        
        return jsonify({
            'success': True,
            'message': f'Usuario temporal creado exitosamente. Expira en {duration_minutes} minutos.',
            'expires_at': expires_at.strftime('%d/%m/%Y %H:%M:%S'),
            'login_token': login_token
        })
        
    except Exception as e:
        return jsonify({'error': f'Error al crear usuario: {str(e)}'}), 500

@app.route('/admin/edit_user/<int:user_id>', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    
    user = User.query.get_or_404(user_id)
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        role = request.form.get('role', 'user')
        
        # Verificar si el nombre de usuario ya existe (excluyendo el usuario actual)
        existing_user = User.query.filter_by(username=username).first()
        if existing_user and existing_user.id != user.id:
            flash('El nombre de usuario ya existe.', 'error')
        else:
            user.username = username
            user.role = role
            
            # Solo actualizar la contraseña si se proporciona una nueva
            if password:
                user.password_hash = generate_password_hash(password)
            
            db.session.commit()
            flash('Usuario actualizado exitosamente.', 'success')
            return redirect(url_for('admin_users'))
    
    return render_template('edit_user.html', user=user)

@app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
@login_required
def delete_user(user_id):
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    
    user = User.query.get_or_404(user_id)
    
    # No permitir eliminar el usuario actual
    if user.id == current_user.id:
        flash('No puedes eliminar tu propia cuenta.', 'error')
        return redirect(url_for('admin_users'))
    
    db.session.delete(user)
    db.session.commit()
    flash('Usuario eliminado exitosamente.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/login', methods=['GET', 'POST'])
@login_required
def admin_login():
    # Verificar si el usuario ya está autenticado como admin
    if session.get('admin_authenticated'):
        # Verificar si la sesión de admin no ha expirado (30 minutos)
        admin_login_time = session.get('admin_login_time')
        if admin_login_time:
            login_time = datetime.fromisoformat(admin_login_time)
            if datetime.now(timezone.utc) - login_time < timedelta(minutes=30):
                return redirect(url_for('admin_panel'))
            else:
                # Sesión expirada, limpiar
                session.pop('admin_authenticated', None)
                session.pop('admin_login_time', None)
    
    # Verificar si el usuario tiene rol de admin - redirigir silenciosamente si no es admin
    if current_user.role != 'admin':
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        password = request.form.get('password')
        
        # Verificar la contraseña del usuario actual (como GitHub)
        if check_password_hash(current_user.password_hash, password):
            session['admin_authenticated'] = True
            session['admin_login_time'] = datetime.now(timezone.utc).isoformat()
            # Configurar la sesión para que expire cuando se cierre el navegador
            session.permanent = False
            flash('Autenticación de administrador exitosa.', 'success')
            return redirect(url_for('admin_panel'))
        else:
            flash('Contraseña incorrecta.', 'error')
    
    return render_template('admin_login.html')

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin_authenticated', None)
    flash('Sesión de administrador cerrada.', 'info')
    return redirect(url_for('dashboard'))

@app.route('/admin/delete_temporary_user/<int:user_id>', methods=['POST'])
@login_required
def delete_temporary_user(user_id):
    # Verificar si el usuario es administrador
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    
    # Verificar si el usuario está autenticado como admin
    if not check_admin_session():
        return jsonify({'error': 'No autenticado como administrador'}), 401
    
    try:
        user = User.query.get_or_404(user_id)
        
        # Verificar que sea un usuario temporal
        if not user.is_temporary:
            return jsonify({'error': 'Solo se pueden eliminar usuarios temporales'}), 400
        
        db.session.delete(user)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Usuario temporal {user.username} eliminado exitosamente.'
        })
        
    except Exception as e:
        return jsonify({'error': f'Error al eliminar usuario: {str(e)}'}), 500

@app.route('/admin/get_temporary_user_qr/<int:user_id>')
@login_required
def get_temporary_user_qr(user_id):
    if not check_admin_session():
        return jsonify({'error': 'No autenticado como administrador'}), 401
        
    user = User.query.get_or_404(user_id)
    
    if not user.is_temporary or not user.login_token:
        return jsonify({'error': 'El usuario no es temporal o no tiene un token de acceso.'}), 400
        
    # Crear la URL de acceso directo con el token
    login_url = url_for('qr_login', token=user.login_token, _external=True)
    
    return jsonify({
        'success': True,
        'qr_data': login_url  # Devolver directamente la URL para el QR
    })

@app.route('/qr_login/<token>')
def qr_login(token):
    # Buscar el usuario por el token, que no esté usado y no haya expirado
    user = User.query.filter_by(login_token=token, login_token_used=False).first()

    if user and user.expires_at > datetime.now(timezone.utc):
        # Marcar el token como usado para que no se pueda volver a usar
        user.login_token_used = True
        db.session.commit()
        
        # Iniciar sesión del usuario
        login_user(user)
        
        flash(f'¡Bienvenido, {user.username}! Has iniciado sesión exitosamente.', 'success')
        return redirect(url_for('dashboard'))
    else:
        # Si el token no es válido, está usado o ha expirado
        flash('El código QR es inválido, ya ha sido utilizado o ha expirado.', 'error')
        return redirect(url_for('login'))

@app.route('/admin/get_user_password/<int:user_id>')
@login_required
def get_user_password(user_id):
    if not check_admin_session():
        return jsonify({'error': 'Acceso denegado'}), 403
    
    user = User.query.get_or_404(user_id)
    
    if not user.is_temporary or not user.temp_password:
        return jsonify({'error': 'No es un usuario temporal o no tiene contraseña temporal'}), 404
        
    return jsonify({'success': True, 'password': user.temp_password})

@app.route('/admin/get_links/<int:year>')
@login_required
def get_links(year):
    if not check_admin_session():
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        # Obtener todas las planillas del año especificado
        planillas = Planilla.query.filter_by(año=year).all()
        
        # Organizar enlaces por tipo y mes
        links = {
            'salud': {},
            'laboratorio': {}
        }
        
        for planilla in planillas:
            # Determinar el tipo basado en el título o descripción
            if 'salud' in planilla.titulo.lower() or 'casanare' in planilla.titulo.lower():
                links['salud'][planilla.mes] = planilla.url_google_drive
            elif 'laboratorio' in planilla.titulo.lower() or 'lab' in planilla.titulo.lower():
                links['laboratorio'][planilla.mes] = planilla.url_google_drive
        
        return jsonify({
            'success': True,
            'links': links
        })
        
    except Exception as e:
        return jsonify({'error': f'Error al obtener enlaces: {str(e)}'}), 500

@app.route('/admin/save_links', methods=['POST'])
@login_required
def save_links():
    print("=== INICIO DE SAVE_LINKS ===")
    print(f"Usuario actual: {current_user.username}")
    print(f"Rol del usuario: {current_user.role}")
    
    if not check_admin_session():
        print("ERROR: No hay sesión de admin activa")
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        print("Recibiendo datos JSON...")
        data = request.get_json()
        print(f"Datos recibidos: {data}")
        
        year = data.get('year')
        links = data.get('links', {})
        
        print(f"Año: {year}")
        print(f"Enlaces: {links}")
        
        if not year:
            print("ERROR: Año no especificado")
            return jsonify({'error': 'Año no especificado'}), 400
        
        # Procesar enlaces de Salud Casanare
        print("Procesando enlaces de Salud Casanare...")
        for mes in range(1, 13):
            url = links.get('salud', {}).get(str(mes)) or links.get('salud', {}).get(mes)
            planilla = Planilla.query.filter_by(mes=mes, año=year).filter(Planilla.titulo.ilike('%salud%')).first()
            if url and url.strip():
                print(f"Procesando mes {mes} para Salud Casanare: {url}")
                if planilla:
                    planilla.titulo = f'Planilla de Control - Salud Casanare - {mes}/{year}'
                    planilla.descripcion = f'Planilla de control de servicios de aseo para Salud Casanare - {mes}/{year}'
                    planilla.url_google_drive = url
                    print(f"Planilla existente actualizada: {planilla.titulo}")
                else:
                    nombres_meses = {
                        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
                        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
                        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
                    }
                    nombre_mes = nombres_meses.get(mes, f'Mes {mes}')
                    planilla = Planilla(
                        mes=mes,
                        año=year,
                        titulo=f'Planilla de Control - Salud Casanare - {nombre_mes} {year}',
                        descripcion=f'Planilla de control de servicios de aseo para Salud Casanare - {nombre_mes} {year}',
                        url_google_drive=url,
                        created_by=current_user.id
                    )
                    db.session.add(planilla)
                    print(f"Nueva planilla creada: {planilla.titulo}")
            else:
                # Si el campo está vacío y existe la planilla, eliminarla
                if planilla:
                    db.session.delete(planilla)
                    print(f"Planilla de Salud Casanare eliminada para mes {mes} año {year}")
        
        # Procesar enlaces de Laboratorio
        print("Procesando enlaces de Laboratorio...")
        for mes in range(1, 13):
            url = links.get('laboratorio', {}).get(str(mes)) or links.get('laboratorio', {}).get(mes)
            planilla = Planilla.query.filter_by(mes=mes, año=year).filter(Planilla.titulo.ilike('%laboratorio%')).first()
            if url and url.strip():
                print(f"Procesando mes {mes} para Laboratorio: {url}")
                if planilla:
                    planilla.titulo = f'Planilla de Control - Laboratorio - {mes}/{year}'
                    planilla.descripcion = f'Planilla de control de servicios de aseo para Laboratorio - {mes}/{year}'
                    planilla.url_google_drive = url
                    print(f"Planilla existente actualizada: {planilla.titulo}")
                else:
                    nombres_meses = {
                        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
                        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
                        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
                    }
                    nombre_mes = nombres_meses.get(mes, f'Mes {mes}')
                    planilla = Planilla(
                        mes=mes,
                        año=year,
                        titulo=f'Planilla de Control - Laboratorio - {nombre_mes} {year}',
                        descripcion=f'Planilla de control de servicios de aseo para Laboratorio - {nombre_mes} {year}',
                        url_google_drive=url,
                        created_by=current_user.id
                    )
                    db.session.add(planilla)
                    print(f"Nueva planilla creada: {planilla.titulo}")
            else:
                # Si el campo está vacío y existe la planilla, eliminarla
                if planilla:
                    db.session.delete(planilla)
                    print(f"Planilla de Laboratorio eliminada para mes {mes} año {year}")
        
        print("Guardando en la base de datos...")
        db.session.commit()
        print("Base de datos actualizada exitosamente")
        
        response_data = {
            'success': True,
            'message': f'Enlaces guardados exitosamente para el año {year}'
        }
        print(f"Respuesta: {response_data}")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"ERROR en save_links: {str(e)}")
        db.session.rollback()
        return jsonify({'error': f'Error al guardar enlaces: {str(e)}'}), 500

@app.route('/admin/update_global_year', methods=['POST'])
@login_required
def update_global_year():
    if not check_admin_session():
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        data = request.get_json()
        new_year = data.get('year')
        
        if not new_year:
            return jsonify({'error': 'Año no especificado'}), 400
        
        # Guardar el año en la sesión
        session['global_year'] = new_year
        
        # Actualizar títulos de planillas existentes para el año anterior
        old_year = session.get('global_year', 2026)  # Usar el año actual de la sesión
        planillas_to_update = Planilla.query.filter_by(año=old_year).all()
        
        nombres_meses = {
            1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
            5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
            9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
        }
        
        for planilla in planillas_to_update:
            nombre_mes = nombres_meses.get(planilla.mes, f'Mes {planilla.mes}')
            
            # Determinar el tipo de planilla basado en el título
            if 'salud' in planilla.titulo.lower() or 'casanare' in planilla.titulo.lower():
                planilla.titulo = f'Planilla de Control - Salud Casanare - {nombre_mes} {new_year}'
                planilla.descripcion = f'Planilla de control de servicios de aseo para Salud Casanare - {nombre_mes} {new_year}'
            elif 'laboratorio' in planilla.titulo.lower() or 'lab' in planilla.titulo.lower():
                planilla.titulo = f'Planilla de Control - Laboratorio - {nombre_mes} {new_year}'
                planilla.descripcion = f'Planilla de control de servicios de aseo para Laboratorio - {nombre_mes} {new_year}'
            
            # Actualizar el año de la planilla
            planilla.año = new_year
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Año global actualizado a {new_year}. Se actualizaron {len(planillas_to_update)} planillas.',
            'year': new_year,
            'updated_planillas': len(planillas_to_update)
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Error al actualizar año: {str(e)}'}), 500

@app.route('/descargar_planilla_2026/<int:mes>/<formato>')
@login_required
def descargar_planilla_2026(mes, formato):
    # Obtener el año dinámico
    year_2026 = session.get('global_year', 2026)
    
    # Obtener la planilla del mes específico para el año dinámico
    planilla = Planilla.query.filter_by(mes=mes, año=year_2026).first()
    
    if not planilla:
        flash('Planilla no encontrada', 'error')
        return redirect(url_for('dashboard'))
    
    # Nombres de los meses en español
    nombres_meses = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    nombre_mes = nombres_meses.get(mes, f'Mes_{mes}')
    nombre_archivo = f"Planilla_Salud_Casanare_{nombre_mes}_{year_2026}"
    
    try:
        # Extraer el ID del archivo de la URL de Google Drive
        url_pattern = r'/d/([a-zA-Z0-9-_]+)/'
        match = re.search(url_pattern, planilla.url_google_drive)
        
        if not match:
            flash('URL de Google Drive no válida', 'error')
            return redirect(url_for('salud_casanare2026_google', mes=mes))
        
        file_id = match.group(1)
        
        if formato == 'excel':
            # URL de descarga directa para Excel
            download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"
            response = requests.get(download_url)
            
            if response.status_code == 200:
                import tempfile
                import os
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
                temp_file.write(response.content)
                temp_file.close()
                
                return send_from_directory(
                    os.path.dirname(temp_file.name),
                    os.path.basename(temp_file.name),
                    as_attachment=True,
                    download_name=f"{nombre_archivo}.xlsx"
                )
            else:
                flash('Error al descargar el archivo Excel', 'error')
                return redirect(url_for('salud_casanare2026_google', mes=mes))
                
        elif formato == 'pdf':
            # URL de descarga directa para PDF
            download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=pdf"
            response = requests.get(download_url)
            
            if response.status_code == 200:
                import tempfile
                import os
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                temp_file.write(response.content)
                temp_file.close()
                
                return send_from_directory(
                    os.path.dirname(temp_file.name),
                    os.path.basename(temp_file.name),
                    as_attachment=True,
                    download_name=f"{nombre_archivo}.pdf"
                )
            else:
                flash('Error al descargar el archivo PDF', 'error')
                return redirect(url_for('salud_casanare2026_google', mes=mes))
                
        elif formato == 'word':
            # Descargar la planilla real de Google Drive como PDF y crear Word elegante
            try:
                # Descargar como PDF desde Google Drive para mantener formato
                pdf_download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=pdf"
                response = requests.get(pdf_download_url)
                
                if response.status_code == 200:
                    import tempfile
                    import os
                    from docx import Document
                    from docx.shared import Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.shared import OxmlElement, qn
                    
                    # Crear documento Word elegante
                    doc = Document()
                    
                    # Configurar márgenes
                    sections = doc.sections
                    for section in sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                    
                    # Título principal
                    title = doc.add_heading(f'Planilla de Control - {nombre_mes} {year_2026}', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Línea separadora
                    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Información de la empresa
                    company_heading = doc.add_heading('Grupo Servis Aseo S.L', level=1)
                    company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    services_para = doc.add_paragraph('Servicios de Limpieza y Mantenimiento')
                    services_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Detalles en formato tabla
                    doc.add_paragraph('')  # Espacio
                    
                    # Crear tabla de información
                    info_table = doc.add_table(rows=3, cols=2)
                    info_table.style = 'Table Grid'
                    info_table.autofit = True
                    
                    # Información de la tabla
                    info_table.cell(0, 0).text = 'Fecha de Descarga:'
                    info_table.cell(0, 1).text = datetime.now(timezone.utc).strftime("%d/%m/%Y")
                    info_table.cell(1, 0).text = 'Mes:'
                    info_table.cell(1, 1).text = nombre_mes
                    info_table.cell(2, 0).text = 'Año:'
                    info_table.cell(2, 1).text = str(year_2026)
                    
                    # Espacio
                    doc.add_paragraph('')
                    
                    # Nota importante
                    note_heading = doc.add_heading('Información Importante', level=2)
                    note_para = doc.add_paragraph('Esta es una versión en formato Word de la planilla original de Google Drive.')
                    note_para2 = doc.add_paragraph('Para acceder a la planilla completa con formato original, utilice la opción "Descargar PDF".')
                    
                    # Espacio
                    doc.add_paragraph('')
                    
                    # Información de acceso
                    access_heading = doc.add_heading('Acceso a la Planilla Original', level=2)
                    access_para = doc.add_paragraph('• La planilla original se encuentra en Google Drive')
                    access_para2 = doc.add_paragraph('• Para editar o ver en tiempo real, acceda directamente a Google Drive')
                    access_para3 = doc.add_paragraph('• Los datos se actualizan automáticamente en la planilla original')
                    
                    # Pie de página
                    doc.add_paragraph('')
                    footer_para = doc.add_paragraph('Documento generado automáticamente por el sistema de gestión de planillas')
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    footer_para.style = 'Intense Quote'
                    
                    # Guardar archivo Word
                    word_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    word_temp.close()
                    doc.save(word_temp.name)
                    
                    return send_from_directory(
                        os.path.dirname(word_temp.name),
                        os.path.basename(word_temp.name),
                        as_attachment=True,
                        download_name=f"{nombre_archivo}.docx"
                    )
                else:
                    flash('Error al acceder a la planilla de Google Drive', 'error')
                    return redirect(url_for('salud_casanare2026_google', mes=mes))
                    
            except Exception as e:
                flash(f'Error al generar el documento Word: {str(e)}', 'error')
                return redirect(url_for('salud_casanare2026_google', mes=mes))
            
        else:
            flash('Formato no válido', 'error')
            return redirect(url_for('salud_casanare2026_google', mes=mes))
            
    except Exception as e:
        flash(f'Error al descargar el archivo: {str(e)}', 'error')
        return redirect(url_for('salud_casanare2026_google', mes=mes))

# === TOKEN SECRETO PARA CRON ===
CRON_SECRET_TOKEN = os.environ.get('CRON_SECRET_TOKEN', 'mi_token_super_secreto_1234567890')

# === CONFIGURACIÓN KEEP-ALIVE MEJORADA ===
KEEP_ALIVE_INTERVAL = int(os.environ.get('KEEP_ALIVE_INTERVAL', 300))  # 5 minutos por defecto
KEEP_ALIVE_ENABLED = os.environ.get('KEEP_ALIVE_ENABLED', 'True').lower() == 'true'

def get_app_url():
    """Obtiene la URL de la aplicación de forma inteligente"""
    # 1. Intentar desde variable de entorno específica
    url = os.environ.get('KEEP_ALIVE_URL')
    if url:
        return url.rstrip('/')
    
    # 2. Intentar desde RENDER_EXTERNAL_URL
    url = os.environ.get('RENDER_EXTERNAL_URL')
    if url:
        return url.rstrip('/')
    
    # 3. Intentar desde RENDER_EXTERNAL_HOSTNAME
    hostname = os.environ.get('RENDER_EXTERNAL_HOSTNAME')
    if hostname:
        return f"https://{hostname}"
    
    # 4. Intentar construir desde RENDER_SERVICE_NAME
    service_name = os.environ.get('RENDER_SERVICE_NAME')
    if service_name:
        return f"https://{service_name}.onrender.com"
    
    # 5. Fallback: usar localhost para desarrollo
    port = os.environ.get('PORT', 5000)
    return f"http://localhost:{port}"

def keep_alive_improved():
    """Sistema de keep-alive mejorado con logging y manejo de errores"""
    app_url = get_app_url()
    
    if not app_url:
        print("⚠️  KEEP-ALIVE: No se pudo determinar la URL de la aplicación")
        return
    
    print(f"🚀 KEEP-ALIVE: Iniciando con URL: {app_url}")
    print(f"⏰ KEEP-ALIVE: Intervalo: {KEEP_ALIVE_INTERVAL} segundos")
    
    consecutive_failures = 0
    max_failures = 5
    
    # Lista de endpoints para probar
    endpoints = ['/ping', '/health', '/', '/status']
    
    while True:
        success = False
        
        # Probar múltiples endpoints
        for endpoint in endpoints:
            try:
                response = requests.get(f"{app_url}{endpoint}", timeout=15)
                
                if response.status_code == 200:
                    if consecutive_failures > 0:
                        print(f"✅ KEEP-ALIVE: Conexión restaurada después de {consecutive_failures} fallos (endpoint: {endpoint})")
                        consecutive_failures = 0
                    else:
                        print(f"✅ KEEP-ALIVE: Ping exitoso - {endpoint} - Status: {response.status_code}")
                    success = True
                    break  # Si un endpoint funciona, no probar los demás
                    
            except requests.exceptions.Timeout:
                print(f"⏰ KEEP-ALIVE: Timeout en {endpoint}")
                continue
                
            except requests.exceptions.ConnectionError:
                print(f"🔌 KEEP-ALIVE: Error de conexión en {endpoint}")
                continue
                
            except Exception as e:
                print(f"❌ KEEP-ALIVE: Error en {endpoint}: {str(e)}")
                continue
        
        if not success:
            consecutive_failures += 1
            print(f"❌ KEEP-ALIVE: Todos los endpoints fallaron (intento {consecutive_failures})")
        
        # Si hay muchos fallos consecutivos, aumentar el intervalo
        if consecutive_failures >= max_failures:
            print(f"🚨 KEEP-ALIVE: Demasiados fallos consecutivos ({consecutive_failures}). Pausando...")
            time.sleep(KEEP_ALIVE_INTERVAL * 2)  # Doble tiempo
        else:
            time.sleep(KEEP_ALIVE_INTERVAL)

@app.route('/cron/ejecutar-tarea')
def cron_ejecutar_tarea():
    token = request.args.get('token')
    if token != CRON_SECRET_TOKEN:
        return 'No autorizado', 403
    try:
        cleanup_expired_users()
        return 'Tarea ejecutada correctamente', 200
    except Exception as e:
        return f'Error al ejecutar la tarea: {str(e)}', 500

# === RUTA DE HEALTH CHECK ===
@app.route('/health')
def health_check():
    """Endpoint de health check para monitoreo externo"""
    try:
        # Verificar conexión a base de datos
        db.session.execute(text('SELECT 1'))
        
        # Verificar usuarios activos
        active_users = User.query.filter(
            User.is_temporary == False
        ).count()
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'database': 'connected',
            'active_users': active_users,
            'uptime': 'running'
        }), 200
        
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'error': str(e)
        }), 500

# === RUTA DE PING SIMPLE ===
@app.route('/ping')
def ping():
    """Endpoint simple de ping para keep-alive"""
    return jsonify({
        'pong': True,
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'message': 'Aplicación activa'
    }), 200

# === RUTA DE STATUS DETALLADO ===
@app.route('/status')
def status():
    """Endpoint de status detallado para administradores"""
    try:
        # Estadísticas del sistema
        total_users = User.query.count()
        active_temporary_users = User.query.filter(
            User.is_temporary == True,
            User.expires_at > datetime.now(timezone.utc)
        ).count()
        
        # Verificar keep-alive
        keep_alive_status = "enabled" if KEEP_ALIVE_ENABLED else "disabled"
        
        return jsonify({
            'status': 'operational',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'system_stats': {
                'total_users': total_users,
                'active_temporary_users': active_temporary_users,
                'keep_alive': keep_alive_status,
                'keep_alive_interval': KEEP_ALIVE_INTERVAL
            },
            'environment': {
                'render': bool(os.environ.get('RENDER')),
                'port': os.environ.get('PORT', 5000),
                'python_version': sys.version
            }
        }), 200
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'error': str(e)
        }), 500

@app.route('/admin/important_message', methods=['GET', 'POST'])
@login_required
def important_message_admin():
    if current_user.role != 'admin':
        return jsonify({'error': 'Acceso denegado. Solo los administradores pueden acceder a esta sección.'}), 403
    if not check_admin_session():
        return redirect(url_for('admin_login'))
    msg = ImportantMessage.query.first()
    if request.method == 'POST':
        if request.is_json:
            data = request.get_json()
            content = data.get('content', '').strip()
            is_active = bool(data.get('is_active'))
        else:
            content = request.form.get('content', '').strip()
            is_active = bool(request.form.get('is_active'))
        if msg:
            msg.content = content
            msg.is_active = is_active
            msg.updated_at = datetime.now(timezone.utc)
        else:
            msg = ImportantMessage(content=content, is_active=is_active)
            db.session.add(msg)
        db.session.commit()
        if request.is_json:
            return jsonify({'success': True, 'message': 'Mensaje actualizado correctamente.'})
        flash('Mensaje actualizado correctamente.', 'success')
        return redirect(url_for('admin_panel'))
    # Para AJAX, devolver datos JSON si se pide
    if request.args.get('json') == '1':
        return jsonify({
            'content': msg.content if msg else '',
            'is_active': msg.is_active if msg else False
        })
    return render_template('admin_important_message.html', msg=msg)

@app.route('/get_important_message')
def get_important_message():
    msg = ImportantMessage.query.first()
    if msg and msg.is_active and msg.content.strip():
        return jsonify({'active': True, 'content': msg.content})
    return jsonify({'active': False})

if __name__ == '__main__':
    import os
    import threading
    import time
    import requests

    with app.app_context():
        init_db()
        # Iniciar el hilo de limpieza en segundo plano
        from threading import Thread

        def background_cleanup():
            while True:
                cleanup_expired_users()
                time.sleep(600)  # 10 minutos

        cleanup_thread = Thread(target=background_cleanup)
        cleanup_thread.daemon = True
        cleanup_thread.start()

        # --- Hilo keep-alive mejorado ---
        def start_keep_alive():
            if KEEP_ALIVE_ENABLED:
                print("🚀 Iniciando sistema de keep-alive...")
                keep_alive_improved()
            else:
                print("⏸️  Keep-alive deshabilitado por configuración")

        # Solo iniciar el keep-alive si estamos en Render o está habilitado
        if (os.environ.get('RENDER') or 
            os.environ.get('RENDER_EXTERNAL_URL') or 
            os.environ.get('RENDER_EXTERNAL_HOSTNAME') or
            KEEP_ALIVE_ENABLED):
            
            ka_thread = threading.Thread(target=start_keep_alive)
            ka_thread.daemon = True
            ka_thread.start()
            print("✅ Hilo de keep-alive iniciado")
        else:
            print("ℹ️  Keep-alive no iniciado (no estamos en Render)")

    port = int(os.environ.get('PORT', 5000))
    # Solo usar debug en desarrollo local, no en producción
    debug_mode = os.environ.get('FLASK_ENV') == 'development' or os.environ.get('DEBUG') == 'True'
    app.run(debug=debug_mode, host='0.0.0.0', port=port) 